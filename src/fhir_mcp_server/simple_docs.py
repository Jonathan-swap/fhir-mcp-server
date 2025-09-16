# Copyright (c) 2025, WSO2 LLC. (https://www.wso2.com/) All Rights Reserved.

# WSO2 LLC. licenses this file to you under the Apache License,
# Version 2.0 (the "License"); you may not use this file except
# in compliance with the License.
# You may obtain a copy of the License at

# http://www.apache.org/licenses/LICENSE-2.0

# Unless required by applicable law or agreed to in writing,
# software distributed under the License is distributed on an
# "AS IS" BASIS, WITHOUT WARRANTIES OR CONDITIONS OF ANY
# KIND, either express or implied. See the License for the
# specific language governing permissions and limitations
# under the License.

import logging
from typing import List, Dict, Any
from pathlib import Path

from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

# Global simple document store
_doc_store = None

def get_doc_store():
    """Get or create the document store."""
    global _doc_store
    if _doc_store is None:
        _doc_store = SimpleDocStore()
    return _doc_store

class SimpleDocStore:
    """Minimal document storage and retrieval."""
    
    def __init__(self):
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=200
        )
        self.vector_store = Chroma(
            persist_directory="./docs_db",
            embedding_function=self.embeddings,
            collection_name="fhir_docs"
        )
    
    def add_text(self, text: str, title: str = "Document") -> None:
        """Add text to the store."""
        doc = Document(page_content=text, metadata={"title": title})
        chunks = self.text_splitter.split_documents([doc])
        self.vector_store.add_documents(chunks)
        self.vector_store.persist()
        logger.info(f"Added document: {title}")
    
    def add_docs(self, paths: List[str | Path]) -> None:
        """Minimal file ingester that scrapes text from .pdf, .docx, and .md files and stores it like add_text."""
        for p in paths:
            path = Path(p)
            if not path.exists() or not path.is_file():
                logger.warning(f"Skipping non-file path: {path}")
                continue

            suffix = path.suffix.lower()
            text = ""

            if suffix == ".md":
                try:
                    text = path.read_text(encoding="utf-8", errors="ignore")
                except Exception as exc:
                    logger.error(f"Failed to read markdown file {path}: {exc}")
                    continue
            elif suffix == ".docx":
                try:
                    from docx import Document as DocxDocument  # type: ignore
                    doc = DocxDocument(str(path))
                    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                except Exception as exc:
                    logger.error(
                        f"Failed to extract text from DOCX {path}. Install python-docx. Error: {exc}"
                    )
                    continue
            elif suffix == ".pdf":
                try:
                    import pypdf  # type: ignore
                    reader = pypdf.PdfReader(str(path))
                    text = "\n".join((page.extract_text() or "") for page in reader.pages)
                except Exception as exc:
                    logger.error(
                        f"Failed to extract text from PDF {path}. Install pypdf. Error: {exc}"
                    )
                    continue
            else:
                logger.warning(f"Unsupported file type for {path}. Only .pdf, .docx, .md are handled.")
                continue

            if text.strip():
                self.add_text(text=text, title=path.name)
            else:
                logger.warning(f"No extractable text found in {path}")

    def search(self, query: str, k: int = 5) -> List[Dict[str, Any]]:
        """Search documents."""
        results = self.vector_store.similarity_search_with_score(query, k=k)
        return [
            {
                "content": doc.page_content,
                "title": doc.metadata.get("title", "Untitled"),
                "score": float(score)
            }
            for doc, score in results
        ]
